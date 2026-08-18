#!/usr/bin/env python3
"""
generate_docx.py — Build an eCH-standard .docx from content.yaml.

Chapters and appendices may reference .adoc files via:
  type: adoc_file
  file: docs/chapters/01-einleitung.adoc

The adoc_parser module converts supported AsciiDoc constructs to block dicts.
All other blocks are rendered directly from the YAML.
"""

import argparse
import os
import shutil
import sys
import yaml
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, os.path.dirname(__file__))
from adoc_parser import parse_adoc


# ── Low-level XML helpers ──────────────────────────────────────────────────────

def make_sdt_text(tag_val, alias_val, value_text):
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    alias = OxmlElement('w:alias'); alias.set(qn('w:val'), alias_val)
    tag   = OxmlElement('w:tag');   tag.set(qn('w:val'), tag_val)
    lock  = OxmlElement('w:lock');  lock.set(qn('w:val'), 'sdtLocked')
    for child in (alias, tag, lock):
        sdtPr.append(child)
    sdt.append(sdtPr)
    sdtContent = OxmlElement('w:sdtContent')
    run = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = value_text
    run.append(t); sdtContent.append(run); sdt.append(sdtContent)
    return sdt


def make_sdt_date(value_text):
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    lock  = OxmlElement('w:lock'); lock.set(qn('w:val'), 'sdtLocked')
    date  = OxmlElement('w:date'); date.set(qn('w:fullDate'), '2024-01-01T00:00:00Z')
    fmt   = OxmlElement('w:dateFormat'); fmt.set(qn('w:val'), 'yyyy-MM-dd')
    lid   = OxmlElement('w:lid');   lid.set(qn('w:val'), 'de-CH')
    store = OxmlElement('w:storeMappedDataAs'); store.set(qn('w:val'), 'dateTime')
    cal   = OxmlElement('w:calendar'); cal.set(qn('w:val'), 'gregorian')
    for child in (fmt, lid, store, cal):
        date.append(child)
    for child in (lock, date):
        sdtPr.append(child)
    sdt.append(sdtPr)
    sdtContent = OxmlElement('w:sdtContent')
    run = OxmlElement('w:r')
    t = OxmlElement('w:t'); t.text = value_text
    run.append(t); sdtContent.append(run); sdt.append(sdtContent)
    return sdt


def make_ref_field(bookmark_name, display_text):
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), f' REF  {bookmark_name} ')
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    alias = OxmlElement('w:alias')
    alias.set(qn('w:val'), 'Rechte Maustaste - Feld aktualisieren')
    tag = OxmlElement('w:tag'); tag.set(qn('w:val'), bookmark_name)
    lock = OxmlElement('w:lock'); lock.set(qn('w:val'), 'sdtLocked')
    sdtPr.append(alias); sdtPr.append(tag); sdtPr.append(lock)
    sdt.append(sdtPr)
    sdtContent = OxmlElement('w:sdtContent')
    run = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = display_text
    run.append(t); sdtContent.append(run); sdt.append(sdtContent)
    fld.append(sdt)
    return fld


# ── Document structure helpers ─────────────────────────────────────────────────

def clear_body(doc):
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def add_toc(doc, title='Inhaltsverzeichnis'):
    p_title = doc.add_paragraph(style='TitelInhaltsverzeichnis')
    p_title.add_run(title)
    p = doc.add_paragraph()
    try:
        p.style = doc.styles['Verzeichnis1']
    except KeyError:
        pass
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r' TOC \o "1-5" \h \z \u ')
    run = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '[Inhaltsverzeichnis wird beim Öffnen in Word aktualisiert]'
    run.append(t); fld.append(run); p._p.append(fld)


def add_tof(doc, heading_text, field_name):
    p = doc.add_paragraph(style='Anhangberschrift')
    p.add_run(heading_text)
    p2 = doc.add_paragraph()
    try:
        p2.style = doc.styles['Abbildungsverzeichnis']
    except KeyError:
        pass
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), rf' TOC \h \z \c "{field_name}" ')
    run = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = f'[{field_name}-Verzeichnis wird beim Öffnen in Word aktualisiert]'
    run.append(t); fld.append(run); p2._p.append(fld)


def add_caption(doc, caption_type, counter, label):
    counter[caption_type] = counter.get(caption_type, 0) + 1
    p = doc.add_paragraph(style='Beschriftung')
    r1 = OxmlElement('w:r')
    t1 = OxmlElement('w:t')
    t1.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t1.text = f'{caption_type} '; r1.append(t1); p._p.append(r1)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), f' SEQ {caption_type} \\* ARABIC ')
    r_f = OxmlElement('w:r')
    t_f = OxmlElement('w:t'); t_f.text = str(counter[caption_type])
    r_f.append(t_f); fld.append(r_f); p._p.append(fld)
    r2 = OxmlElement('w:r')
    t2 = OxmlElement('w:t')
    t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t2.text = f': {label}'; r2.append(t2); p._p.append(r2)


def add_generic_table(doc, headers, rows, col_widths=None):
    n_cols = max(len(headers) if headers else 0,
                 max((len(r) for r in rows), default=0))
    if n_cols == 0:
        return
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = 'Tabellenraster'
    if col_widths:
        tbl = table._tbl
        tblGrid = OxmlElement('w:tblGrid')
        for cw in col_widths:
            gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(cw))
            tblGrid.append(gc)
        tbl_pr = tbl.find(qn('w:tblPr'))
        if tbl_pr is not None:
            tbl_pr.addnext(tblGrid)
    if headers:
        hrow = table.add_row()
        for i, h in enumerate(headers):
            p = hrow.cells[i].paragraphs[0]
            p.style = doc.styles['Tabellentitel']
            p.add_run(h)
    for data_row in rows:
        drow = table.add_row()
        for i, val in enumerate(data_row[:n_cols]):
            p = drow.cells[i].paragraphs[0]
            p.style = doc.styles['Tabellentext']
            p.add_run(str(val))


def add_image(doc, image_path, width_cm=6.0, align='right'):
    if not os.path.exists(image_path):
        doc.add_paragraph(f'[Bild nicht gefunden: {image_path}]', style='Standard')
        return
    p = doc.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if align == 'right'
                   else WD_ALIGN_PARAGRAPH.LEFT)
    p.add_run().add_picture(image_path, width=Cm(width_cm))


def add_runs_to_para(p, runs_or_text):
    """Add text or run list to an existing paragraph."""
    if isinstance(runs_or_text, str):
        p.add_run(runs_or_text)
    else:
        for r in runs_or_text:
            run = p.add_run(r.get('text', ''))
            run.bold   = r.get('bold', False)
            run.italic = r.get('italic', False)


# ── Block renderer ────────────────────────────────────────────────────────────

def render_blocks(doc, blocks, images_dir, caption_counter):
    """Render a list of block dicts into the document."""
    for block in blocks:
        btype = block.get('type', 'paragraph')

        # ── adoc_file: parse and recurse ──────────────────────
        if btype == 'adoc_file':
            filepath = block.get('file', '')
            if not os.path.exists(filepath):
                doc.add_paragraph(f'[adoc_file not found: {filepath}]',
                                  style='Standard')
                continue
            sub_blocks = parse_adoc(filepath, images_dir=images_dir)
            render_blocks(doc, sub_blocks, images_dir, caption_counter)

        # ── Headings ──────────────────────────────────────────
        elif btype == 'heading':
            style_map = {1: 'berschrift1', 2: 'berschrift2', 3: 'berschrift3',
                         4: 'berschrift4', 5: 'berschrift5'}
            p = doc.add_paragraph(
                style=style_map.get(block.get('level', 1), 'berschrift1'))
            p.add_run(block.get('text', ''))

        elif btype == 'anhang_heading':
            doc.add_paragraph(style='Anhangberschrift').add_run(block.get('text', ''))

        elif btype == 'nebentitel':
            doc.add_paragraph(style='Nebentitel').add_run(block.get('text', ''))

        elif btype == 'subtitle':
            doc.add_paragraph(style='Untertitel').add_run(block.get('text', ''))

        elif btype == 'austauschformat':
            doc.add_paragraph(style='Austauschformat').add_run(block.get('text', ''))

        # ── Paragraph (plain or with inline runs) ─────────────
        elif btype == 'paragraph':
            style = block.get('style', 'Standard')
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style='Standard')
            runs = block.get('runs')
            if runs:
                add_runs_to_para(p, runs)
            else:
                r = p.add_run(block.get('text', ''))
                r.bold   = block.get('bold', False)
                r.italic = block.get('italic', False)

        # ── List items ────────────────────────────────────────
        elif btype == 'list_item':
            style_map = {
                'bullet1': 'Aufzhlung',
                'bullet2': 'Liste-',
                'bullet3': 'Liste-',
                'number1': 'Liste1',
                'alpha1':  'Listea',
            }
            style = style_map.get(block.get('list_style', 'bullet1'), 'Aufzhlung')
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style='Aufzhlung')
            runs = block.get('runs')
            if runs:
                add_runs_to_para(p, runs)
            else:
                p.add_run(block.get('text', ''))

        # ── Image ─────────────────────────────────────────────
        elif btype == 'image':
            add_image(doc,
                      os.path.join(images_dir, block.get('file', '')),
                      width_cm=block.get('width_cm', 6.0),
                      align=block.get('align', 'right'))
            if 'caption' in block:
                add_caption(doc, 'Abbildung', caption_counter, block['caption'])

        # ── Table ─────────────────────────────────────────────
        elif btype == 'table':
            add_generic_table(doc,
                              block.get('headers', []),
                              block.get('rows', []),
                              block.get('col_widths'))
            if 'caption' in block:
                add_caption(doc, 'Tabelle', caption_counter, block['caption'])

        elif btype == 'empty':
            doc.add_paragraph(style='Standard')


# ── Cover page ────────────────────────────────────────────────────────────────

def render_cover(doc, metadata):
    ech_nr = metadata.get('ech_nummer', '<ID>')
    title  = metadata.get('title', '<Name>')

    # Title paragraph with bookmarked SDTs
    p_title = doc.add_paragraph(style='Titel')
    pel = p_title._p

    def bm_start(bm_id, name):
        el = OxmlElement('w:bookmarkStart')
        el.set(qn('w:id'), str(bm_id)); el.set(qn('w:name'), name); return el

    def bm_end(bm_id):
        el = OxmlElement('w:bookmarkEnd')
        el.set(qn('w:id'), str(bm_id)); return el

    pel.append(bm_start(100, 'eCHNummer'))
    pel.append(make_sdt_text('eCH-Dossier Nr.', 'eCH-Dossier Nr.', ech_nr))
    pel.append(bm_end(100))
    r_sep = OxmlElement('w:r')
    t_sep = OxmlElement('w:t')
    t_sep.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_sep.text = ' – '; r_sep.append(t_sep); pel.append(r_sep)
    pel.append(bm_start(101, 'eCHName'))
    pel.append(make_sdt_text('eCH-Dossier Name', 'eCH-Dossier Name', title))
    pel.append(bm_end(101))

    # Metadata table
    _render_metadata_table(doc, metadata, ech_nr, title)

    # Zusammenfassung
    doc.add_paragraph(style='Standard')
    doc.add_paragraph(style='Nebentitel').add_run('Zusammenfassung')
    p_zus = doc.add_paragraph(style='Standard')
    p_zus.add_run(metadata.get('zusammenfassung',
                  '<Kurze Zusammenfassung des Zwecks des Dokuments>')).italic = True
    doc.add_paragraph(style='Nebentitel')  # spacer


def _render_metadata_table(doc, metadata, ech_nr, title):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Tabellenraster'
    tbl = table._tbl
    tblGrid = OxmlElement('w:tblGrid')
    for cw in [2622, 6589]:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(cw))
        tblGrid.append(gc)
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is not None:
        tbl_pr.addnext(tblGrid)

    def add_row(label, value_fn):
        row = table.add_row()
        p_l = row.cells[0].paragraphs[0]
        p_l.style = doc.styles['Tabellentitel']; p_l.add_run(label)
        p_v = row.cells[1].paragraphs[0]
        p_v.style = doc.styles['Tabellentext']
        value_fn(p_v, row.cells[1])

    m = metadata

    def ref(bm, display):
        return lambda p, c: p._p.append(make_ref_field(bm, display))

    def plain(text):
        return lambda p, c: p.add_run(text)

    def date_sdt(text):
        return lambda p, c: p._p.append(make_sdt_date(text))

    def ersetzt(version, status):
        def _fn(p, c):
            p._p.append(make_sdt_text('Version', 'Version', version))
            r = OxmlElement('w:r'); t = OxmlElement('w:t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = ' – '; r.append(t); p._p.append(r)
            p._p.append(make_sdt_text('Antragsart', 'Antragsart', status))
        return _fn

    def herausgeber(p, c):
        lines = [m.get('herausgeber', 'Verein eCH, Affolternstrasse 52, 8050 Zürich'),
                 m.get('kontakt', 'T 044 388 74 64 / info@ech.ch / www.ech.ch')]
        p.add_run(lines[0])
        for line in lines[1:]:
            p2 = c.add_paragraph(style='Tabellentext'); p2.add_run(line)

    add_row('Name',                ref('eCHName', title))
    add_row('eCH-Nummer',          ref('eCHNummer', m.get('ech_nummer', '<ID>')))
    add_row('Kategorie',           plain(m.get('kategorie', 'Standard')))
    add_row('Reifegrad',           plain(m.get('reifegrad', 'Stabil')))
    add_row('Version',             plain(m.get('version', 'x.x.x')))
    add_row('Status',              plain(m.get('status', 'Entwurf')))
    add_row('Beschluss am',        date_sdt(m.get('beschluss', 'JJJJ-MM-TT')))
    add_row('Ausgabedatum',        date_sdt(m.get('ausgabe', 'JJJJ-MM-TT')))
    add_row('Ersetzt Version',     ersetzt(m.get('ersetzt_version', 'x.x.x'),
                                           m.get('ersetzt_status', 'Aufgehoben')))
    add_row('Voraussetzungen',     plain(m.get('voraussetzungen', '')))
    add_row('Beilagen',            plain(m.get('beilagen', '')))
    add_row('Sprachen',            plain(m.get('sprachen', '')))
    add_row('Fachgruppe',          plain(m.get('fachgruppe', '')))
    add_row('Herausgeber / Vertrieb', herausgeber)


# ── Hinweis ───────────────────────────────────────────────────────────────────

def render_hinweis(doc):
    doc.add_paragraph(style='Nebentitel').add_run('Hinweis')
    doc.add_paragraph(style='Standard').add_run(
        'Im vorliegenden Dokument wird bei der Bezeichnung von Personen eine '
        'geschlechtsneutrale Formulierung verwendet. Basis bildet der Leitfaden '
        'der Bundeskanzlei. Je nach Situation kommen Paarformen (Bürgerinnen und '
        'Bürger), geschlechtsabstrakte Formen (versicherte Person), '
        'geschlechtsneutrale Formen (Versicherte) oder Umschreibungen ohne '
        'Personenbezug zum Einsatz. Das generische Maskulin (Bürger) ist nicht '
        'zulässig. Vollformen werden in fortlaufenden Texten verwendet. In '
        'verknappten Textpassagen, namentlich in Tabellen, können Kurzformen '
        'verwendet werden (Referent/in). Genderstern und ähnliche Schreibweisen '
        'werden nicht verwendet.'
    )


# ── Appendices ────────────────────────────────────────────────────────────────

def render_appendices(doc, appendices, images_dir, caption_counter):
    labels = list('ABCDEFGH')
    for i, app in enumerate(appendices):
        label = labels[i] if i < len(labels) else str(i + 1)
        heading = f'Anhang {label} – {app.get("title", "")}'
        app_type = app.get('type', 'generic')

        if app_type == 'abbildungsverzeichnis':
            add_tof(doc, heading, 'Abbildung')
        elif app_type == 'tabellenverzeichnis':
            add_tof(doc, heading, 'Tabelle')
        else:
            doc.add_paragraph(style='Anhangberschrift').add_run(heading)
            adoc_file = app.get('adoc_file')
            if adoc_file:
                render_blocks(doc, [{'type': 'adoc_file', 'file': adoc_file}],
                              images_dir, caption_counter)
            else:
                render_blocks(doc, app.get('blocks', []), images_dir, caption_counter)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--reference', required=True)
    parser.add_argument('--content',   required=True)
    parser.add_argument('--output',    required=True)
    parser.add_argument('--images',    default='docs/images')
    args = parser.parse_args()

    with open(args.content, encoding='utf-8') as f:
        content = yaml.safe_load(f)

    shutil.copy2(args.reference, args.output)
    doc = Document(args.output)
    clear_body(doc)

    metadata        = content.get('metadata', {})
    caption_counter = {}

    render_cover(doc, metadata)
    add_toc(doc)
    render_hinweis(doc)
    render_blocks(doc, content.get('chapters', []), args.images, caption_counter)
    render_appendices(doc, content.get('appendices', []), args.images, caption_counter)

    doc.save(args.output)
    print(f'Generated: {args.output}')


if __name__ == '__main__':
    main()
